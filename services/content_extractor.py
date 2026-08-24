"""Extracción de texto a partir de lo que aporte el profesor: texto pegado, PDF, Word o una URL puntual."""
import re

import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document

USER_AGENT = "Mozilla/5.0 (compatible; HerramientaDocente/1.0; uso educativo personal)"


class ExtractionError(RuntimeError):
    pass


def extract_from_pdf(file_stream) -> str:
    try:
        with pdfplumber.open(file_stream) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:  # noqa: BLE001 - queremos un mensaje claro para el usuario
        raise ExtractionError(f"No se pudo leer el PDF: {exc}") from exc
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text:
        raise ExtractionError("No se ha encontrado texto en el PDF (¿es un escaneado sin OCR?).")
    return text


def extract_from_docx(file_stream) -> str:
    try:
        document = Document(file_stream)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"No se pudo leer el documento Word: {exc}") from exc
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    text = "\n".join(parts)
    if not text.strip():
        raise ExtractionError("El documento Word no contiene texto.")
    return text


def extract_from_url(url: str) -> str:
    if not re.match(r"^https?://", url.strip(), re.IGNORECASE):
        raise ExtractionError("La URL debe empezar por http:// o https://")
    try:
        resp = requests.get(url, headers={"User-Agent": USER_AGENT}, timeout=15)
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise ExtractionError(f"No se pudo descargar la URL indicada: {exc}") from exc

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "aside"]):
        tag.decompose()

    main = soup.find("article") or soup.find("main") or soup.body or soup
    text = main.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    cleaned = "\n".join(lines)
    if not cleaned:
        raise ExtractionError("No se ha podido extraer texto legible de esa URL.")
    return cleaned


def extract_text(source_type: str, *, text: str = "", file_storage=None, url: str = "") -> str:
    """Punto de entrada único para las rutas Flask.

    source_type: "texto" | "pdf" | "docx" | "url"
    """
    if source_type == "texto":
        if not text.strip():
            raise ExtractionError("Pega algo de texto antes de continuar.")
        return text.strip()

    if source_type == "pdf":
        if file_storage is None:
            raise ExtractionError("Sube un archivo PDF.")
        return extract_from_pdf(file_storage)

    if source_type == "docx":
        if file_storage is None:
            raise ExtractionError("Sube un archivo Word (.docx).")
        return extract_from_docx(file_storage)

    if source_type == "url":
        if not url.strip():
            raise ExtractionError("Indica una URL.")
        return extract_from_url(url.strip())

    raise ExtractionError(f"Origen de contenido no soportado: {source_type}")
